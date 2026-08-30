"""
Build the presenter's prep sheet for the 13-slide deck.

One block per slide: what is on it, what it means, the sentence to say,
and the question to expect. Written to be read in 10 to 20 minutes.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = os.environ.get(
    "PREP_OUT",
    r"g:\codes\Ass\DTMS\submission\Presentation_Prep_Notes.docx")

INK   = RGBColor(0x1A, 0x1D, 0x21)
RED   = RGBColor(0xA8, 0x32, 0x2A)
BLUE  = RGBColor(0x1F, 0x4E, 0x5F)
GREY  = RGBColor(0x60, 0x65, 0x6A)

doc = Document()

# page setup
for sec in doc.sections:
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)

normal = doc.styles["Normal"]
normal.font.name = "Segoe UI"
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(4)


def para(text="", size=10.5, bold=False, color=INK, italic=False,
         space_before=0, space_after=4, indent=0.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    r.font.name = "Segoe UI"
    return p


def rich(parts, size=10.5, space_before=0, space_after=4, indent=0.0):
    """parts: list of (text, bold, color)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    for t, b, c in parts:
        r = p.add_run(t)
        r.font.size = Pt(size)
        r.font.bold = b
        r.font.color.rgb = c
        r.font.name = "Segoe UI"
    return p


def slide_block(n, title, shows, means, say, ask=None, answer=None):
    rich([("SLIDE %d" % n, True, RED), ("   ", False, INK),
          (title, True, INK)], size=12, space_before=12, space_after=3)
    rich([("What is on it.  ", True, GREY), (shows, False, INK)], indent=0.15)
    rich([("What it means.  ", True, GREY), (means, False, INK)], indent=0.15)
    rich([("Say this.  ", True, BLUE), (say, False, INK)], indent=0.15)
    if ask:
        rich([("If asked: ", True, RED), (ask, True, INK), ("  ", False, INK),
              (answer, False, INK)], indent=0.15, space_after=2)


# ---------------------------------------------------------------- header
para("Presentation prep sheet", size=19, bold=True, space_after=2)
para("DTMS_Paper_Presentation_Short.pptx  |  13 slides  |  12 minutes  |  "
     "about 55 seconds per slide", size=10, color=GREY, space_after=10)

rich([("The one-sentence version of the whole talk. ", True, INK),
      ("We measured whether common database tuning advice actually works, "
       "on two databases and 23,480 measurements, and all four pieces of "
       "advice we tested make things worse under conditions we can now name.",
       False, INK)], size=11, space_after=8)

rich([("The one thing to keep straight. ", True, RED),
      ("Slides 1 to 7 compare two databases, PostgreSQL and MariaDB. "
       "Slides 8, 9 and 10 are about PostgreSQL only. The slide says so at "
       "the top. If you remember nothing else, remember that, because it is "
       "the question you are most likely to be asked.", False, INK)],
     size=11, space_after=6)

doc.add_paragraph()

# ---------------------------------------------------------------- part 1
para("PART 1  |  SETUP  (slides 1-5, about 4 minutes)", size=11, bold=True,
     color=RED, space_before=6, space_after=4)

slide_block(
    1, "Title",
    "Title, your name, supervisor, and the scale of the study.",
    "Nothing to explain.",
    "\"This is a study of how databases decide which index to use, and "
    "whether the advice people give about it actually works.\"")

slide_block(
    2, "Motivation: a decision nobody sees",
    "The one decision every query needs, and the four pieces of advice we "
    "tested.",
    "Every query needs one choice made for it: scan the whole table, or use "
    "an index, and if so which one. The user never sees it, but it can make "
    "the query a hundred times slower.",
    "\"Because this decision is invisible, a lot of advice has grown around "
    "it. It sounds sensible, and almost none of it has been measured. That "
    "is what I did.\"")

slide_block(
    3, "Why two databases",
    "The reason the study uses PostgreSQL and MariaDB rather than one.",
    "If you test one database and find a problem, you cannot tell whether "
    "you found something true about databases in general or just a bug in "
    "that one product. Two independently written databases let you separate "
    "those.",
    "\"If both databases fail the same way, it is a design problem. If only "
    "one fails, it is that product's choice. Every result later is labelled "
    "one way or the other.\"")

slide_block(
    4, "The two measures",
    "Definitions of regret and q-error.",
    "Regret = how many times slower the plan it chose was, compared to the "
    "best plan it could have chosen. Regret of 1.0 means it did the best it "
    "could. q-error = how badly it guessed the number of rows. 1.0 is a "
    "perfect guess.",
    "\"I keep these two apart on purpose, because the main finding is that "
    "they come apart: the guess was right and the choice was still wrong.\"",
    "Why not just measure speed?",
    "Because then a slow database looks bad even when it made the right "
    "choice. Regret measures the decision, not the engine.")

slide_block(
    5, "The experiment",
    "The table of how many measurements at each table size.",
    "23,480 measurements. Both databases cover the same seven table sizes, "
    "so the comparison is fair.",
    "\"I generated the data instead of using a standard benchmark, because "
    "that way I know the true row count for every query. So I can say the "
    "database guessed wrong, instead of guessing along with it.\"")

# ---------------------------------------------------------------- part 2
para("PART 2  |  COMPARING THE TWO DATABASES  (slides 6-7, about 2 minutes)",
     size=11, bold=True, color=RED, space_before=12, space_after=4)

slide_block(
    6, "Selection quality: PostgreSQL 2.3%, MariaDB 66.0%",
    "A two-row table comparing how often each database chose badly.",
    "PostgreSQL picks the best available plan 97.7% of the time. MariaDB "
    "gets it wrong on two thirds of queries, but its mistakes are milder "
    "(median 1.34x versus PostgreSQL's worst of 1.39x).",
    "\"Read the last two columns together. PostgreSQL is almost always "
    "right. MariaDB is usually wrong, but only by a little. Two different "
    "kinds of failure, not one database being better at everything.\"",
    "Is the 2.3% with or without BRIN?",
    "Without. That is exactly what slide 8 is about, so say \"I am coming to "
    "that in two slides.\"")

slide_block(
    7, "Why they differ: the plans available",
    "A count of what each database actually chose: bitmap scan, index "
    "lookup, or full table scan.",
    "PostgreSQL uses a bitmap scan 346 times; MariaDB 7 times. MariaDB gave "
    "up and scanned the whole table 66 times; PostgreSQL 9 times.",
    "\"A bitmap scan sorts the matching rows into disk order before fetching "
    "them. That puts a ceiling on how bad a mistake can be, because even a "
    "wrong choice still reads the disk in order. MariaDB has no equivalent, "
    "so it has no ceiling.\"",
    "So MariaDB is just worse?",
    "Not exactly. It is wrong more often but more mildly. Slide 11 shows "
    "that reverses as tables grow.")

# ---------------------------------------------------------------- part 3
para("PART 3  |  THE MAIN FINDING, IN THREE STEPS  "
     "(slides 8-10, about 3 minutes)",
     size=11, bold=True, color=RED, space_before=12, space_after=3)
rich([("These three slides are one story, in order: ", False, INK),
      ("what happened", True, INK), (" \u2192 ", False, GREY),
      ("what it was not", True, INK), (" \u2192 ", False, GREY),
      ("what it actually was", True, INK),
      (".  All three are PostgreSQL only.", False, INK)],
     size=10.5, space_after=6)

slide_block(
    8, "Adding one more index broke it",
    "Two rows. Without a BRIN index: 2.3% bad choices. With one: 73.3%.",
    "Same database, same data, same queries. The only thing that changed is "
    "whether a BRIN index exists on the column next to the B-tree. People "
    "add BRIN indexes because they are tiny and look free.",
    "\"On slide 6 I said PostgreSQL chooses well, 2.3% bad. That was without "
    "a BRIN index. Watch what happens when I add one: 73.3%. This is not a "
    "contradiction of slide 6, it is what adding one supposedly free index "
    "does to it.\"",
    "Is this both databases?",
    "No, PostgreSQL only. MariaDB does not have BRIN indexes at all, so it "
    "cannot make this mistake. That is how we know it is a PostgreSQL "
    "costing defect and not something unavoidable.")

slide_block(
    9, "It was not bad estimates",
    "98.2% of those 113 bad choices had the row count essentially right.",
    "The usual explanation for a bad query plan is that the database "
    "misjudged how many rows would come back. Here it did not. It knew the "
    "row count, it knew both indexes existed, and it still chose the slower "
    "one.",
    "\"This rules out the usual fix too. Running ANALYZE or adding "
    "statistics cannot help, because you cannot improve a number that was "
    "already correct. So if the estimates were right, what did it get "
    "wrong?\"",
    "Where does 113 come from?",
    "It is the bad choices from the previous slide: PostgreSQL at 1,000,000 "
    "rows. 110 of the 113 come from the with-BRIN case, which is why these "
    "two slides belong together.")

slide_block(
    10, "It compares the wrong thing",
    "BRIN costs 12.13 to read and takes 65.95 ms. B-tree costs 213.35 to "
    "read and takes 0.34 ms.",
    "When two indexes can answer the same query, PostgreSQL keeps whichever "
    "is cheaper to read and throws the other away. But reading the index is "
    "only half the job; the query then has to fetch the actual rows, and "
    "that part is not counted. BRIN is tiny so it always wins, and the index "
    "thrown away is the fast one.",
    "\"The index that costs 17 times less to read is the one that runs 194 "
    "times slower. It is comparing the wrong quantity.\"",
    "How do you know this is the cause?",
    "I read the PostgreSQL source. The function is choose_bitmap_and in "
    "indxpath.c, and it compares index-read cost only. It is on the slide.")

# ---------------------------------------------------------------- part 4
para("PART 4  |  THE WIDER PATTERN  (slides 11-13, about 3 minutes)",
     size=11, bold=True, color=RED, space_before=12, space_after=4)

slide_block(
    11, "The two databases fail in opposite directions",
    "Two charts: how often each database errs, and how badly, as tables grow.",
    "As tables get bigger, PostgreSQL is wrong more often (0% up to 36%) but "
    "never by more than 3.7x. MariaDB is wrong less often (53% down to 12%) "
    "but its worst case climbs to 30.9x. They cross over.",
    "\"A database that is wrong often but never badly is easier to run than "
    "one wrong rarely but catastrophically. If you summarise either with a "
    "single number you lose exactly this.\"")

slide_block(
    12, "Four practices, all four backfire",
    "Four rows: what each piece of advice promises, and what we measured.",
    "In three of the four, the advice does exactly what it claims and the "
    "query still gets slower. Extended statistics really do fix the estimate "
    "by 108x. The plan gets worse anyway.",
    "\"That is the pattern of the whole study. The advice improves the thing "
    "it promises to improve, and that thing is not what determines how fast "
    "the query runs.\"")

slide_block(
    13, "What to do, and what not to do",
    "Four things to do on the left, four not to do on the right, each with "
    "the number behind it.",
    "This is the practical payoff. Every 'do not' has a measured number "
    "attached so it is a finding, not an opinion.",
    "\"All four were reasonable advice. None of it shipped with the "
    "conditions under which it stops being true. That is what this study "
    "adds: the conditions.\"")

# ---------------------------------------------------------------- back page
doc.add_page_break()
para("The four numbers you must not mix up", size=14, bold=True,
     space_after=6)

rows = [
    ("2.3%", "PostgreSQL, WITHOUT a BRIN index. Its normal, good behaviour. "
             "Slides 6 and 8."),
    ("73.3%", "PostgreSQL, WITH a BRIN index added. Same data, same queries. "
              "Slide 8."),
    ("98.2%", "Of the 113 cases where PostgreSQL chose badly, the share that "
              "had a correct row estimate. Slide 9."),
    ("194x", "How much slower one query got, purely because a BRIN index "
             "existed. Slide 10."),
    ("66.0%", "MariaDB's bad-choice rate. Nothing to do with BRIN; MariaDB "
              "has no BRIN. Slide 6."),
]
for num, meaning in rows:
    rich([(num, True, RED), ("   ", False, INK), (meaning, False, INK)],
         size=11, space_after=6)

para("Why 2.3% and 73.3% are not a contradiction", size=12, bold=True,
     space_before=10, space_after=4)
para("They are the same database on the same data. The only difference is "
     "which indexes existed when the query ran. Without BRIN it chooses well. "
     "Add a BRIN index beside the B-tree and it chooses badly, because of the "
     "costing defect on slide 10. If someone thinks you contradicted "
     "yourself, that one sentence resolves it.", size=10.5)

para("Three questions worth rehearsing", size=12, bold=True,
     space_before=12, space_after=4)
qa = [
    ("Why did you not use a standard benchmark like TPC-H?",
     "Because a fixed dataset fixes the skew, the correlation and the "
     "clustering at whatever values it happens to have, so you cannot "
     "attribute any effect to one property. Generating the data also means I "
     "know the true row count for every query."),
    ("How do you know the results are not just noise?",
     "Every query runs seven times, the first is discarded, and I report the "
     "median with its spread. Nineteen automated checks run before any "
     "measurement is accepted, including validating my generated data "
     "against PostgreSQL's own internal statistics."),
    ("Did you find anything you got wrong yourself?",
     "Yes, three measurement errors, and they are all reported in the paper "
     "rather than removed. One of them was an explanation I had to withdraw "
     "after testing my own claim and finding it did not hold."),
]
for q, a in qa:
    rich([(q, True, INK)], size=10.5, space_before=6, space_after=2)
    para(a, size=10.5, indent=0.2, space_after=2)

para("If you run short of time", size=12, bold=True, space_before=12,
     space_after=4)
para("Slides 8, 9, 10 and 13 are the ones that matter. Slide 7 and slide 11 "
     "can be summarised in one sentence each and skipped. Do not skip slide "
     "13; the recommendations are what the audience takes away.",
     size=10.5)

doc.save(OUT)
print("wrote", OUT)
